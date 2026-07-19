from __future__ import annotations

import re
import sqlite3
import tempfile
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

from docx import Document

QUESTION_PREFIX = re.compile(r"^\s*\**\s*\d+\s*[.)]\s*\**\s*")
ANSWER_MARK = re.compile(r"^\s*[*+#]+\s*")
ANSWER_LABEL = re.compile(r"^\s*[A-ZА-ЯЁ]\s*[.)]\s*", re.IGNORECASE)
SEPARATOR = re.compile(r"^[-–—_\s]+$")
SPACE = re.compile(r"\s+")
TAG = re.compile(r"<[^>]+>")


@dataclass
class ParsedAnswer:
    text: str
    correct: bool = False


@dataclass
class ParsedQuestion:
    question: str
    answers: list[ParsedAnswer]
    valid: bool = True
    problems: list[str] = field(default_factory=list)
    duplicate_in_file: bool = False
    duplicate_in_database: bool = False
    source_name: str | None = None

    def to_dict(self) -> dict[str, Any]:
        return asdict(self)


def clean_text(value: str) -> str:
    value = value.replace("\u00a0", " ").replace("\u200b", " ").replace("\ufeff", " ")
    value = value.replace("**", "")
    return SPACE.sub(" ", value).strip()


def clean_question(value: str) -> str:
    return clean_text(QUESTION_PREFIX.sub("", value))


def clean_answer(value: str, strip_label: bool = True) -> ParsedAnswer:
    raw = clean_text(value)
    correct = bool(ANSWER_MARK.match(raw))
    raw = ANSWER_MARK.sub("", raw)
    if strip_label:
        raw = ANSWER_LABEL.sub("", raw)
    return ParsedAnswer(text=clean_text(raw), correct=correct)


def validate_question(item: ParsedQuestion) -> ParsedQuestion:
    problems: list[str] = []
    if not item.question:
        problems.append("Savol matni bo'sh")
    item.answers = [answer for answer in item.answers if answer.text]
    if len(item.answers) < 2:
        problems.append("Javob variantlari 2 tadan kam")
    correct_count = sum(1 for answer in item.answers if answer.correct)
    if correct_count != 1:
        problems.append(f"To'g'ri javoblar soni {correct_count} ta")
    normalized = [answer.text.casefold() for answer in item.answers]
    if len(normalized) != len(set(normalized)):
        problems.append("Bir xil javob varianti takrorlangan")
    item.problems = problems
    item.valid = not problems
    return item


def mark_file_duplicates(items: list[ParsedQuestion]) -> list[ParsedQuestion]:
    seen: set[str] = set()
    for item in items:
        key = item.question.casefold().strip()
        if key and key in seen:
            item.duplicate_in_file = True
        seen.add(key)
    return items


def _build_from_blocks(blocks: list[list[str]]) -> list[ParsedQuestion]:
    result: list[ParsedQuestion] = []
    for block in blocks:
        lines = [clean_text(line) for line in block if clean_text(line) and not SEPARATOR.match(clean_text(line))]
        if len(lines) < 2:
            continue
        question = clean_question(lines[0])
        strip_labels = all(bool(ANSWER_LABEL.match(ANSWER_MARK.sub("", line))) for line in lines[1:])
        answers = [clean_answer(line, strip_label=strip_labels) for line in lines[1:]]
        result.append(validate_question(ParsedQuestion(question=question, answers=answers)))
    return mark_file_duplicates(result)


def parse_lines(lines: list[str]) -> list[ParsedQuestion]:
    cleaned = [line.rstrip("\r\n") for line in lines]
    has_numbered_questions = sum(1 for line in cleaned if QUESTION_PREFIX.match(line)) >= 1

    if has_numbered_questions:
        blocks: list[list[str]] = []
        current: list[str] = []
        for line in cleaned:
            text = clean_text(line)
            if not text or SEPARATOR.match(text):
                continue
            if QUESTION_PREFIX.match(text):
                if current:
                    blocks.append(current)
                current = [text]
            elif current:
                current.append(text)
        if current:
            blocks.append(current)
        return _build_from_blocks(blocks)

    blocks = []
    current = []
    for line in cleaned:
        text = clean_text(line)
        if not text:
            if current:
                blocks.append(current)
                current = []
            continue
        if not SEPARATOR.match(text):
            current.append(text)
    if current:
        blocks.append(current)
    return _build_from_blocks(blocks)


def parse_txt(content: bytes) -> list[ParsedQuestion]:
    text: str | None = None
    for encoding in ("utf-8-sig", "cp1251", "utf-16"):
        try:
            text = content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        raise ValueError("TXT fayl kodirovkasini aniqlab bo'lmadi")
    return parse_lines(text.splitlines())


def parse_docx(content: bytes) -> list[ParsedQuestion]:
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as handle:
        handle.write(content)
        temp_path = Path(handle.name)
    try:
        document = Document(temp_path)
        if document.tables:
            blocks: list[list[str]] = []
            flattened: list[str] = []
            separate_tables = True
            for table in document.tables:
                rows: list[str] = []
                for row in table.rows:
                    text = clean_text(" ".join(cell.text for cell in row.cells))
                    if text and not SEPARATOR.match(text):
                        rows.append(text)
                        flattened.append(text)
                if rows:
                    if not QUESTION_PREFIX.match(rows[0]):
                        separate_tables = False
                    blocks.append(rows)
            if separate_tables and blocks:
                return _build_from_blocks(blocks)
            return parse_lines(flattened)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if clean_text(paragraph.text)]
        return parse_lines(paragraphs)
    finally:
        temp_path.unlink(missing_ok=True)


def _safe_table_names(connection: sqlite3.Connection) -> set[str]:
    return {row[0] for row in connection.execute("SELECT name FROM sqlite_master WHERE type='table'")}


def parse_db(content: bytes, mode: str = "single") -> list[ParsedQuestion]:
    with tempfile.NamedTemporaryFile(suffix=".db", delete=False) as handle:
        handle.write(content)
        temp_path = Path(handle.name)
    try:
        uri = f"file:{temp_path.as_posix()}?mode=ro"
        connection = sqlite3.connect(uri, uri=True)
        connection.row_factory = sqlite3.Row
        try:
            tables = _safe_table_names(connection)
            required = {"source_questions", "source_answers"}
            missing = required - tables
            if missing:
                raise ValueError(f"Faylda quyidagi jadval topilmadi: {', '.join(sorted(missing))}")

            has_sources = "sources" in tables
            source_map: dict[int, str] = {}
            if has_sources:
                for row in connection.execute("SELECT id, name FROM sources"):
                    source_map[int(row["id"])] = str(row["name"])

            questions = connection.execute(
                "SELECT id, source_id, question_text FROM source_questions ORDER BY id"
            ).fetchall()
            answer_columns = {row[1] for row in connection.execute("PRAGMA table_info(source_answers)")}
            answer_order = "COALESCE(position, id), id" if "position" in answer_columns else "id"
            result: list[ParsedQuestion] = []
            for question_row in questions:
                answer_rows = connection.execute(
                    f"SELECT answer_text, is_correct FROM source_answers WHERE question_id=? ORDER BY {answer_order}",
                    (question_row["id"],),
                ).fetchall()
                answers = [
                    ParsedAnswer(
                        text=clean_text(TAG.sub("", str(answer_row["answer_text"]))),
                        correct=bool(answer_row["is_correct"]),
                    )
                    for answer_row in answer_rows
                ]
                source_name = source_map.get(int(question_row["source_id"])) if mode == "full" else None
                result.append(
                    validate_question(
                        ParsedQuestion(
                            question=clean_text(TAG.sub("", str(question_row["question_text"]))),
                            answers=answers,
                            source_name=source_name,
                        )
                    )
                )
            return mark_file_duplicates(result)
        finally:
            connection.close()
    finally:
        temp_path.unlink(missing_ok=True)


def parse_uploaded_file(filename: str, content: bytes, db_mode: str = "single") -> list[ParsedQuestion]:
    suffix = Path(filename).suffix.lower()
    if suffix == ".txt":
        return parse_txt(content)
    if suffix == ".docx":
        return parse_docx(content)
    if suffix in {".db", ".sqlite", ".sqlite3"}:
        return parse_db(content, mode=db_mode)
    raise ValueError("Faqat .txt, .docx va .db fayllar qabul qilinadi")
